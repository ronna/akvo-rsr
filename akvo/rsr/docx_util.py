# -*- coding: utf-8 -*-
"""Akvo RSR is covered by the GNU Affero General Public License.

See more details in the license.txt file located at the root folder of the Akvo RSR module.
For additional details on the GNU license please see < http://www.gnu.org/licenses/agpl.html >.
"""

import re
import markdown

from bs4 import BeautifulSoup, NavigableString, Tag
from docxtpl import RichText
from docx.opc import constants
from six import text_type, PY2


def markdown_to_docx(text, tpl):
    if not text:
        return ''

    source_code = markdown.markdown(text, do_terms=False)
    source_code = re.sub("\n", ' ', source_code)
    source_code = re.sub(">\s+<", '><', source_code)
    soup = BeautifulSoup('<html>' + source_code + '</html>', 'html.parser')
    parser = SoupParser(tpl)
    for elem in soup.find_all(recursive=False):
        parser.traverse(elem)
    output = text_type(parser)

    return output


class SoupParser(object):

    def __init__(self, tpl):
        self.paragraphs = [dict(params=dict(style='p', indentation=0), runs=[RichText('')])]
        self.current_paragraph = self.paragraphs[-1]
        self.run = self.current_paragraph['runs'][-1]
        self.bold = False
        self.italic = False
        self.underline = False
        self.strike = False
        self.indentation = 0
        self.style = 'p'
        self.still_new = True
        self.size = None
        self.tpl = tpl

    def new_paragraph(self):
        if self.still_new:
            # print("new_paragraph is still new and style is " + self.style + " and indentation is " + text_type(self.indentation))
            self.current_paragraph['params']['style'] = self.style
            self.current_paragraph['params']['indentation'] = self.indentation
            return

        # print("new_paragraph where style is " + self.style + " and indentation is " + text_type(self.indentation))
        self.current_paragraph = dict(params=dict(style=self.style, indentation=self.indentation), runs=[RichText('')])
        self.paragraphs.append(self.current_paragraph)
        self.run = self.current_paragraph['runs'][-1]
        self.still_new = True

    def __str__(self):
        return self.__unicode__().encode('utf-8') if PY2 else self.__unicode__()

    def __unicode__(self):
        output = ''
        list_number = 1
        for para in self.paragraphs:
            # print("Got a paragraph where style is " + para['params']['style'] + " and indentation is " + text_type(para['params']['indentation']))
            output += '<w:p><w:pPr><w:pStyle w:val="Normal"/>'
            if para['params']['style'] in ('ul', 'ol', 'blockquote'):
                output += '<w:ind w:left="' + text_type(36*para['params']['indentation']) + '" w:right="0" w:hanging="0"/>'
            output += '<w:rPr></w:rPr></w:pPr>'
            if para['params']['style'] == 'ul':
                output += text_type(RichText("•\t"))
            if para['params']['style'] == 'ol':
                output += text_type(RichText(text_type(list_number) + ".\t"))
                list_number += 1
            else:
                list_number = 1
            for run in para['runs']:
                output += text_type(run)
            output += '</w:p>'
        return output

    def start_link(self, url):
        ref = self.tpl.docx._part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self.current_paragraph['runs'].append('<w:hyperlink r:id="%s">' % (ref, ))
        self.new_run()
        self.still_new = False

    def end_link(self):
        self.current_paragraph['runs'].append('</w:hyperlink>')
        self.new_run()
        self.still_new = False

    def new_run(self):
        self.current_paragraph['runs'].append(RichText(''))
        self.run = self.current_paragraph['runs'][-1]

    def traverse(self, elem):
        for part in elem.contents:
            if isinstance(part, NavigableString):
                self.run.add(text_type(part), italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size)
                self.still_new = False
            elif isinstance(part, Tag):
                # print("Part name is " + text_type(part.name))
                if part.name == 'p':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'li':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'ul':
                    # print("Entering a UL")
                    oldstyle = self.style
                    self.style = 'ul'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    # print("Leaving a UL")
                elif part.name == 'ol':
                    # print("Entering a OL")
                    oldstyle = self.style
                    self.style = 'ol'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    # print("Leaving a OL")
                elif part.name == 'strong':
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                elif part.name == 'em':
                    self.italic = True
                    self.traverse(part)
                    self.italic = False
                elif part.name == 'strike':
                    self.strike = True
                    self.traverse(part)
                    self.strike = False
                elif part.name == 'u':
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                elif part.name == 'blockquote':
                    oldstyle = self.style
                    self.style = 'blockquote'
                    self.indentation += 20
                    self.traverse(part)
                    self.indentation -= 20
                    self.style = oldstyle
                elif re.match(r'h[1-6]', part.name):
                    oldsize = self.size
                    self.size = 60 - ((int(part.name[1]) - 1) * 10)
                    self.new_paragraph()
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                    self.size = oldsize
                elif part.name == 'a':
                    self.start_link(part['href'])
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                    self.end_link()
                elif part.name == 'br':
                    self.run.add("\n", italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size)
                    self.still_new = False
            else:
                print("Encountered a " + part.__class__.__name__)
